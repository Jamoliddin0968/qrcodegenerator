import os
import qrcode
from io import BytesIO
from django.conf import settings
from django.shortcuts import render
from django.core.files.storage import FileSystemStorage
from docx import Document
from docx.shared import Inches,Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from .models import UploadedFile


def upload_docx(request):
    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']
        fs = FileSystemStorage()
        temp_name = fs.save(uploaded_file.name, uploaded_file)
        temp_path = fs.path(temp_name)

        db_file = UploadedFile.objects.create(
            original_name=uploaded_file.name,
            file=f'uploads/temp_{uploaded_file.name}'
        )

        new_filename = f"{db_file.uuid_name}.docx"
        new_path = os.path.join(settings.MEDIA_ROOT, 'uploads', new_filename)
        os.makedirs(os.path.dirname(new_path), exist_ok=True)

        with open(temp_path, 'rb') as src, open(new_path, 'wb') as dst:
            dst.write(src.read())

        domain = request.build_absolute_uri('/')[:-1]
        verify_url = f"{domain}/verify/{db_file.uuid_name}/"

        qr = qrcode.QRCode(box_size=12, border=2)
        qr.add_data(verify_url)
        qr.make(fit=True)
        buf = BytesIO()
        img = qr.make_image(fill_color="black", back_color="white").convert("RGB")
        img.save(buf, format='PNG', dpi=(300,300))
        # # QR code yaratamiz (fayl emas, verify sahifasiga)
        # qr_img = qrcode.make(verify_url)
        # O()
        # qr_img.save(buf, format='PNG')
        buf.seek(0)

        # DOCX ga QR va kod yozamiz
        doc = Document(new_path)
        table = doc.add_table(rows=1, cols=3)
        # table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        text_cell = table.rows[0].cells[0]
        text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p_text = text_cell.paragraphs[0]
        p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_text = p_text.add_run("This document is a copy of an electronic document generated in accordance with the provision on the Single Portal of Interactive Public Services, approved by the provision of the Cabinet of Ministers of the Republic of Uzbekistan dated September 15, 2017 No. 728. To check the accuracy of the information specified in the copy of the electronic document, go to the website repo.gov.uz and enter the unique number of the electronic document, or scan the QR code using a mobile device. Attention! In accordance with the provision of the Cabinet of Ministers of the Republic of Uzbekistan dated September 15, 2017 No. 728, the information contained in electronic documents is legitimate. It is strictly forbidden for state bodies to refuse to accept copies of electronic documents generated on the Single Portal of Interactive Public Services.")
        run_text.font.size = Pt(10)
        run_text.font.name = 'Times New Roman'
        text_cell.width = Inches(5.09)
        # Kod (chapda)
        cell_code = table.rows[0].cells[1]
        cell_code.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p_code = cell_code.paragraphs[0]
        run_code = p_code.add_run(f"{db_file.code}")
        run_code.bold = False
        run_code.font.size = Pt(23)
        run_code.font.name = 'Cambria'
        # kattaroq shrift
        p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_code.width = Inches(0.98)

        # QR (o‘ngda)
        cell_qr = table.rows[0].cells[2]
        cell_qr.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p_qr = cell_qr.paragraphs[0]
        run_qr = p_qr.add_run()
        run_qr.add_picture(buf, width=Inches(1.22))
        p_qr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell_qr.width = Inches(1.22)
        doc.save(new_path)

        db_file.file.name = f"uploads/{new_filename}"
        db_file.save()

        os.remove(temp_path)

        return render(request, 'result.html', {
            'file_url': f"{settings.MEDIA_URL}uploads/{new_filename}",
            'uuid': db_file.uuid_name,
            'code': db_file.code,
            'verify_url': verify_url
        })

    return render(request, 'upload.html')



from django.shortcuts import get_object_or_404, render
from django.http import FileResponse

def verify_file(request, uuid):
    file_obj = get_object_or_404(UploadedFile, uuid_name=uuid)

    if request.method == 'POST':
        # print(request.POST)
        code = request.POST.get('code')
        if code == file_obj.code:
            file_path = file_obj.file.path
            return FileResponse(open(file_path, 'rb'), as_attachment=True, filename=file_obj.original_name)
        else:
            return render(request, 'index.html', {
                'error': 'Kod noto‘g‘ri!',
                'uuid': uuid
            })

    return render(request, 'index.html', {'uuid': uuid})

import uuid
import qrcode
import base64
import random
from io import BytesIO
from datetime import datetime
from django.http import HttpResponse
from django.template.loader import render_to_string
from weasyprint import HTML
from django.conf import settings
from .models import UploadedFile
import os

import os
import pdfkit
import qrcode
import base64
import uuid
import random
from io import BytesIO
from datetime import datetime
from django.http import HttpResponse
from django.conf import settings
from django.template.loader import render_to_string
from .models import UploadedFile


def create_pdf_view(request):
    if request.method == "POST":
        full_name = request.POST.get("full_name")
        pinfl = request.POST.get("pinfl")

        months = request.POST.getlist("month[]")
        companies = request.POST.getlist("company[]")
        salaries = request.POST.getlist("salary[]")
        taxes = request.POST.getlist("tax[]")

        incomes = []
        for i in range(len(months)):
            if (months[i] or companies[i] or salaries[i] or taxes[i]):
                incomes.append({
                    "month": (months[i] or "").strip(),
                    "company": (companies[i] or "").strip(),
                    "salary": (salaries[i] or "").strip(),
                    "tax": (taxes[i] or "").strip(),
                })

        # создаём запись
        db_file = UploadedFile.objects.create(
            original_name=f"{full_name}.pdf",
            file="",
        )

        # генерим 4-значный код
        code4 = f"{random.randint(0, 9999):04d}"
        try:
            db_file.code = code4
            db_file.save(update_fields=["code"])
        except Exception:
            code4 = db_file.code

        # verify URL + QR
        domain = request.build_absolute_uri('/')[:-1]
        verify_url = f"{domain}/verify/{db_file.uuid_name}/"

        qr_img = qrcode.make(verify_url)
        buf = BytesIO()
        qr_img.save(buf, format='PNG')
        qr_base64 = base64.b64encode(buf.getvalue()).decode('utf-8')
        qr_data = f"data:image/png;base64,{qr_base64}"

        # рендерим HTML
        html_string = render_to_string("pdf_template.html", {
            "full_name": full_name,
            "pinfl": pinfl,
            "date": datetime.now().strftime("%Y-%m-%d"),
            "doc_number": str(uuid.uuid4())[:8],
            "incomes": incomes,
            "verify_url": verify_url,
            "code": code4,
            "qr_data": qr_data,
        })

        # путь для сохранения PDF
        pdf_path = os.path.join(settings.MEDIA_ROOT, "uploads", f"{db_file.uuid_name}.pdf")
        os.makedirs(os.path.dirname(pdf_path), exist_ok=True)

        # === wkhtmltopdf через pdfkit ===
        config = pdfkit.configuration(
            wkhtmltopdf="/usr/bin/wkhtmltopdf"  # путь к бинарнику wkhtmltopdf
        )

        options = {
            'page-size': 'A4',
            'margin-top': '1cm',
            'margin-right': '1cm',
            'margin-bottom': '1cm',
            'margin-left': '1cm',
            'enable-local-file-access': None,  # чтобы разрешить загрузку CSS/изображений
            'encoding': 'UTF-8',
            'quiet': '',
        }

        pdfkit.from_string(
            html_string,
            pdf_path,
            configuration=config,
            options=options
        )

        db_file.file.name = f"uploads/{db_file.uuid_name}.pdf"
        db_file.save(update_fields=["file"])

        return HttpResponse(
            f"<h3>PDF готов:</h3><a href='{settings.MEDIA_URL}uploads/{db_file.uuid_name}.pdf'>Скачать</a>"
        )

    return render(request, "pdf_form.html")
